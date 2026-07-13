#!/usr/bin/env python3
"""Module 5: Seedance JSON, optional views, and machine-integrity final quality."""
from __future__ import annotations

import argparse
import json
import re
import sys
from collections import OrderedDict
from datetime import datetime, timezone
from pathlib import Path

SCRIPTS=Path(__file__).resolve().parent; sys.path.insert(0,str(SCRIPTS))
from core.errors import IntegrityError
from core.execution import execute_node
from core.policy_engine import task_policy
from core.storage import artifact_path, atomic_write_text, file_hashes, load_manifest, read_json, read_jsonl, run_root, sha256_file, write_json
from core.validator import gate_result

TYPE_TO_TEMPLATE={"character":"ROLE","scene":"SCENE","prop":"PROP"}

def _episodes(values):
    result=set()
    for value in values:
        match=re.search(r"(d+)",str(value))
        if match: result.add(int(match.group(1)))
    return sorted(result)

def pre_delivery_audit(manifest_path: Path) -> None:
    gates={aid:read_json(artifact_path(manifest_path,aid,must_exist=True)) for aid in ["source-coverage","recognition-closure","planning-closure","prompt-closure"]}
    source=read_jsonl(artifact_path(manifest_path,"main-image-production-source",must_exist=True)); errors=[]
    for aid,gate in gates.items():
        if gate.get("status")!="passed": errors.append(f"gate failed: {aid}")
    task_ids=[row["task_id"] for row in source]
    if len(task_ids)!=len(set(task_ids)): errors.append("production source contains duplicate task_id")
    if any(row.get("status")!="ready" for row in source): errors.append("production source contains non-ready task")
    checks={aid:gate.get("status") for aid,gate in gates.items()}; checks["production_source_rows"]=len(source)
    result=gate_result("pre-delivery-audit",errors,[],file_hashes(manifest_path,["run-manifest","source-coverage","recognition-closure","planning-closure","prompt-closure","main-image-production-source"]),{},checks=checks)
    write_json(artifact_path(manifest_path,"production-audit"),result)
    if errors: raise IntegrityError("; ".join(errors))

def json_delivery(manifest_path: Path) -> None:
    manifest=load_manifest(manifest_path); audit=read_json(artifact_path(manifest_path,"production-audit",must_exist=True))
    if audit["status"]!="passed": raise IntegrityError("production audit did not pass")
    source=read_jsonl(artifact_path(manifest_path,"main-image-production-source",must_exist=True)); policy=task_policy(); ratios=policy["image_ratios"]
    assets=OrderedDict()
    for row in source:
        base_id=row["base_asset_id"]; template_type=TYPE_TO_TEMPLATE[row["asset_type"]]
        if base_id not in assets:
            assets[base_id]={"assetId":base_id,"templateName":row["asset_name"],"templateType":template_type,"description":row["asset_name"],"appearance":{"description":row["asset_name"],"appearances":[]},"metadata":{"source":"agent","taskIds":[],"evidenceIds":[]}}
        generation={"method":row["task_type"]}
        if row["task_type"]=="image_text_edit": generation["referenceTaskId"]=row["anchor_task_id"]
        appearance={"stateId":row["state_asset_id"] or row["task_id"],"name":row["state_name"],"imageRatio":ratios[template_type],"imagePrompt":row["prompt"] or row["edit_instruction"],"negativePrompt":row["negative_prompt"],"priority":row["priority_level"],"status":"ready","generation":generation,"reviewFocus":row["quality_focus"]}
        episodes=_episodes(row["episode_ids"])
        if episodes: appearance["episodes"]=episodes
        assets[base_id]["appearance"]["appearances"].append(appearance)
        assets[base_id]["metadata"]["taskIds"].append(row["task_id"])
        assets[base_id]["metadata"]["evidenceIds"].extend(row["evidence_ids"])
    for asset in assets.values(): asset["metadata"]["evidenceIds"]=sorted(set(asset["metadata"]["evidenceIds"]))
    output={"schemaVersion":"seedance-element-extract-manifest.v1","manifestVersion":"1.0","generationScope":"main_image_only","source":{"skillId":manifest["skill"]["id"],"skillVersion":manifest["skill"]["version"],"generatedAt":datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds"),"promptMode":"core_prompt_requires_style_prefix","stylePrefixOwner":"middle_platform","stylePrefixRequired":True,"artifactManifest":"run-manifest.json"},"summary":{"totalItems":len(assets),"manualReviewItems":0},"assets":list(assets.values())}
    if manifest["options"].get("content_id"): output["contentId"]=manifest["options"]["content_id"]
    if manifest["options"].get("project_title"): output["project"]={"title":manifest["options"]["project_title"]}
    write_json(artifact_path(manifest_path,"prompt-manifest"),output)

def _markdown_table(rows):
    headers=["task_id","asset_type","asset_name","state_name","priority_level","task_type","episode_ids","scene_ids","evidence_ids","quality_focus"]
    lines=["# 主图生产任务（派生视图）","","> 本文件由 machine/main-image-production-source.jsonl 自动派生，不是机器输入。","","| " + " | ".join(headers) + " |","|"+"|".join("---" for _ in headers)+"|"]
    for row in rows:
        values=[]
        for h in headers:
            value=row[h]; value=",".join(value) if isinstance(value,list) else str(value); values.append(value.replace("|","／").replace("\n","<br>"))
        lines.append("| "+" | ".join(values)+" |")
    return "\n".join(lines)+"\n"

def human_view_delivery(manifest_path: Path) -> None:
    manifest=load_manifest(manifest_path); source_path=artifact_path(manifest_path,"main-image-production-source",must_exist=True); rows=read_jsonl(source_path); generated=[]
    if manifest["options"].get("include_human_views"):
        md_path=run_root(manifest_path)/"views"/"main-image-production-table.md"; atomic_write_text(md_path,_markdown_table(rows)); generated.append(md_path.relative_to(run_root(manifest_path)).as_posix())
        try:
            from docx import Document
            doc=Document(); doc.add_heading("主图生产任务（派生视图）",level=1); table=doc.add_table(rows=1,cols=7); headers=["task_id","asset_type","asset_name","state_name","priority_level","task_type","quality_focus"]
            for i,h in enumerate(headers): table.rows[0].cells[i].text=h
            for row in rows:
                cells=table.add_row().cells
                for i,h in enumerate(headers): cells[i].text=str(row[h])
            docx_path=run_root(manifest_path)/"deliverables"/"main-image-review-table.docx"; docx_path.parent.mkdir(parents=True,exist_ok=True); doc.save(docx_path); generated.append(docx_path.relative_to(run_root(manifest_path)).as_posix())
        except ImportError as exc: raise IntegrityError("python-docx is required when include_human_views=true") from exc
    write_json(artifact_path(manifest_path,"view-manifest"),{"enabled":bool(manifest["options"].get("include_human_views")),"generated_files":generated,"source_artifact":"main-image-production-source","source_sha256":sha256_file(source_path)})

def final_quality(manifest_path: Path) -> None:
    source=read_jsonl(artifact_path(manifest_path,"main-image-production-source",must_exist=True)); prompt=read_json(artifact_path(manifest_path,"prompt-manifest",must_exist=True)); audit=read_json(artifact_path(manifest_path,"production-audit",must_exist=True)); closure=read_json(artifact_path(manifest_path,"prompt-closure",must_exist=True)); read_json(artifact_path(manifest_path,"view-manifest",must_exist=True))
    errors=[]
    if audit["status"]!="passed": errors.append("production audit failed")
    if closure["status"]!="passed": errors.append("prompt closure failed")
    source_tasks={row["task_id"] for row in source}; json_tasks={task_id for asset in prompt.get("assets",[]) for task_id in asset.get("metadata",{}).get("taskIds",[])}
    if source_tasks!=json_tasks: errors.append("production source and prompt manifest task sets differ")
    if prompt.get("summary",{}).get("manualReviewItems")!=0: errors.append("V4 prompt manifest must not contain manual review items")
    result=gate_result("final-quality-gate",errors,[],file_hashes(manifest_path,["run-manifest","main-image-production-source","prompt-manifest","production-audit","prompt-closure","view-manifest"]),{},machine_integrity="invalid" if errors else "valid",delivery_ready=not errors,precision_claim="not_evaluated_at_runtime")
    write_json(artifact_path(manifest_path,"final-quality"),result)
    if errors: raise IntegrityError("; ".join(errors))

NODE_HANDLERS={"pre-delivery-audit":pre_delivery_audit,"json-delivery":json_delivery,"human-view-delivery":human_view_delivery,"final-quality-gate":final_quality}
ACTION_TO_NODE={"pre-delivery-audit":"pre-delivery-audit","json-delivery":"json-delivery","human-view-delivery":"human-view-delivery","final-quality":"final-quality-gate"}
def main(argv=None):
    parser=argparse.ArgumentParser(); parser.add_argument("--run-manifest",required=True,type=Path); parser.add_argument("--action",required=True,choices=sorted(ACTION_TO_NODE)); args=parser.parse_args(argv); node=ACTION_TO_NODE[args.action]; return execute_node(args.run_manifest.resolve(),node,NODE_HANDLERS[node])
if __name__=="__main__": raise SystemExit(main())
